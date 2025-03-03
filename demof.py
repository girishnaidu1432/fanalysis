import streamlit as st
import pandas as pd
import openai
import matplotlib.pyplot as plt
import textwrap
from io import BytesIO
from docx import Document

# Set Streamlit page background
st.markdown(f"""
<style>
    .stApp {{
        background-image: url("https://e0.pxfuel.com/wallpapers/986/360/desktop-wallpaper-background-color-4851-background-color-theme-colorful-brown-color.jpg");
        background-attachment: fixed;
        background-size: cover;
    }}
</style>
""", unsafe_allow_html=True)

# OpenAI API Configuration
openai.api_key = "14560021aaf84772835d76246b53397a"
openai.api_base = "https://amrxgenai.openai.azure.com/"
openai.api_type = 'azure'
openai.api_version = '2024-02-15-preview'
deployment_name = 'gpt'

def analyze_chatbot(question, df):
    prompt = f"""
    Using the data provided below, analyze and respond to the following question:
    {df.to_string(index=False)}
    Question: {question}
    """
    response = openai.ChatCompletion.create(
        engine=deployment_name,
        messages=[{"role": "system", "content": "You are a data analyst expert."},
                  {"role": "user", "content": prompt}],
        temperature=0.7
    )
    return response["choices"][0]["message"]["content"].strip()

def plot_trend(df, group_by_col, value_col, title):
    trend_data = df.groupby(group_by_col)[value_col].sum().reset_index()
    fig, ax = plt.subplots(figsize=(10, 5))  # Increase figure size
    ax.bar(trend_data[group_by_col], trend_data[value_col], color="skyblue")
    plt.xticks(rotation=90, ha='right', fontsize=8)  # Rotate and adjust alignment
    plt.title(title)
    plt.xlabel(group_by_col)
    plt.ylabel(value_col)
    plt.tight_layout()  # Prevent overlapping
    st.markdown("<br><br>", unsafe_allow_html=True)
    st.pyplot(fig)

st.title("Bonus Analysis with OpenAI")
st.write("Upload and analyze CSV data containing bonus-related details.")

uploaded_file = st.file_uploader("Upload CSV File", type=["csv"])
if "search_history" not in st.session_state:
    st.session_state.search_history = []

if uploaded_file:
    try:
        df = pd.read_csv(uploaded_file, encoding="latin1")
        st.write("Uploaded Data Preview:")
        st.dataframe(df.head())
    except UnicodeDecodeError:
        st.error("The uploaded file has an unsupported encoding. Please save it as UTF-8.")
    except Exception as e:
        st.error(f"Error processing the file: {e}")
    else:
        required_columns = {"Partner Id", "Last Name", "Paid As Position", "Gender", "Date of Birth", "Manager Name", "Recruiter Name", "Paid As", "Personal Sales Unit(PSU)", "Team Units(TU)", "First Name", "Adhoc Payment(ADP)", "Recruitment Commission Bonus (RCB)", "Basic commission Bonus(BCB)", "Super Commission Bonus(SCB)", "Performance Bonus (PCB)", "Gross Earnings"}
        if required_columns.issubset(set(df.columns)):
            st.success("File successfully uploaded and validated!")
            
            tab1, tab2 = st.tabs(["Analysis", "Chatbot"])
            
            with tab1:
                st.subheader("Basic Analysis")
                analysis_options = [
                    ("Role-wise Gross Earnings", "Paid As Position", "Gross Earnings"),
                    ("Total Bonus Distribution", "Paid As Position", "Basic commission Bonus(BCB)"),
                    ("Performance Bonus by Position", "Paid As Position", "Performance Bonus (PCB)"),
                    ("Recruitment Commission Analysis", "Recruiter Name", "Recruitment Commission Bonus (RCB)"),
                    ("Manager-wise Bonus Distribution", "Manager Name", "Gross Earnings"),
                    ("Personal Sales Contribution", "First Name", "Personal Sales Unit(PSU)"),
                    ("Team Units Contribution", "First Name", "Team Units(TU)"),
                    ("Adhoc Payments Analysis", "First Name", "Adhoc Payment(ADP)"),
                    ("Gender-Based Earnings", "Gender", "Gross Earnings"),
                    ("Bonus Comparison by Gender", "Gender", "Basic commission Bonus(BCB)")
                ]
                
                for title, group_by, value in analysis_options:
                    st.subheader(title)
                    plot_trend(df, group_by, value, title)
            
            with tab2:
                st.subheader("Chatbot - Insights, Trends, and Analysis")
                user_question = st.text_input("Enter your question:", key="chat_input")
                
                if user_question:
                    response = analyze_chatbot(user_question, df)
                    st.session_state.search_history.append({"question": user_question, "response": response})
                
                if st.session_state.search_history:
                    st.subheader("Search History")
                    for entry in st.session_state.search_history:
                        with st.expander(entry['question']):
                            st.write(entry['response'])
                    
                    if st.button("Download History"):
                        doc = Document()
                        doc.add_heading("Chatbot Search History", level=1)
                        for entry in st.session_state.search_history:
                            doc.add_heading(f"Q: {entry['question']}", level=2)
                            doc.add_paragraph(f"A: {entry['response']}")
                        
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        st.download_button("Download DOCX", buffer, "search_history.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key='download-docx')
        else:
            st.error("Uploaded file is missing required columns.")
